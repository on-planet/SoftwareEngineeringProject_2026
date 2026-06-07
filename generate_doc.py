# -*- coding: utf-8 -*-
"""生成软件测试作业文档（Word + PDF）"""
import os
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_chinese_font(run, font_name='SimSun', size=10.5, bold=False):
    font = run.font
    font.name = font_name
    font.size = Pt(size)
    font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_heading_zh(doc, text, level=1):
    """添加中文标题"""
    if level == 0:
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_chinese_font(run, 'SimHei', 22, True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(18)
        p.paragraph_format.space_before = Pt(12)
        return p
    sizes = {1: 16, 2: 14, 3: 12}
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    set_chinese_font(run, 'SimHei', sizes.get(level, 12), True)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.space_before = Pt(12)
    return p


def add_para_zh(doc, text, bold=False, indent=True, size=10.5):
    """添加中文正文段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_chinese_font(run, 'SimSun', size, bold)
    return p


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_el = OxmlElement(f'w:{edge}')
        edge_el.set(qn('w:val'), 'single')
        edge_el.set(qn('w:sz'), '4')
        edge_el.set(qn('w:space'), '0')
        edge_el.set(qn('w:color'), '000000')
        tcBorders.append(edge_el)
    tcPr.append(tcBorders)


def fill_table_header(table, headers):
    """填充表头并设置样式"""
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr_cells[i].text = text
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_chinese_font(run, 'SimHei', 10.5, True)
        set_cell_border(hdr_cells[i])


def fill_table_row(table, row_idx, values):
    """填充表格行"""
    cells = table.rows[row_idx].cells
    for i, text in enumerate(values):
        cells[i].text = str(text)
        for paragraph in cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_chinese_font(run, 'SimSun', 10.5, False)
        set_cell_border(cells[i])


def build_document():
    doc = Document()

    # 页面边距
    sections = doc.sections[0]
    sections.top_margin = Cm(2.54)
    sections.bottom_margin = Cm(2.54)
    sections.left_margin = Cm(3.17)
    sections.right_margin = Cm(3.17)

    # ===== 封面标题 =====
    add_heading_zh(doc, '第X组-QuantPulse系统-软件测试', level=0)
    doc.add_paragraph()

    # ===== 目录 =====
    add_heading_zh(doc, '目录', level=1)
    toc = [
        '一、测试计划',
        '二、白盒测试（基本路径测试）',
        '    2.1 功能一：期货数据列表查询（list_futures）',
        '    2.2 功能二：烟蒂股策略分析（analyze_cigarbutt）',
        '三、黑盒测试（等价类划分与边界值分析）',
        '    3.1 功能一：期货数据列表查询（list_futures_route）',
        '    3.2 功能二：烟蒂股策略分析（get_cigarbutt_analysis）',
    ]
    for item in toc:
        p = doc.add_paragraph(item, style='List Number' if not item.startswith('    ') else None)
        if item.startswith('    '):
            p.paragraph_format.left_indent = Cm(1.5)
        for run in p.runs:
            set_chinese_font(run, 'SimSun', 12, False)

    # ===== 一、测试计划 =====
    add_heading_zh(doc, '一、测试计划', level=1)

    add_heading_zh(doc, '1.1 项目概述', level=2)
    add_para_zh(doc, '项目名称：QuantPulse 多源金融数据分析平台')
    add_para_zh(doc, '项目简介：QuantPulse 是一个面向 A 股、港股和美股场景的多源金融数据分析平台，采用单仓库（monorepo）结构，包含完整的数据采集（ETL）、REST API 服务（FastAPI）和 Web 前端（Next.js）。系统提供多市场行情、个人空间、组合分析、策略模块、管理后台等核心能力。')
    add_para_zh(doc, '测试工程师：XXX')

    add_heading_zh(doc, '1.2 测试目标', level=2)
    add_para_zh(doc, '（1）验证系统核心功能（期货查询、策略分析、K 线查询等）的正确性与稳定性；')
    add_para_zh(doc, '（2）发现软件中潜在的缺陷，提升软件质量，降低线上故障风险；')
    add_para_zh(doc, '（3）确保系统满足需求规格说明中的功能性与非功能性要求。')

    add_heading_zh(doc, '1.3 测试范围', level=2)
    add_para_zh(doc, '本次测试覆盖以下范围：')
    add_para_zh(doc, '功能测试：期货数据分页查询、烟蒂股策略分析、K 线对比查询、用户管理、组合诊断等；')
    add_para_zh(doc, '接口测试：REST API 接口的请求参数校验、响应格式正确性、异常处理；')
    add_para_zh(doc, '性能测试：核心接口的平均响应时间、并发处理能力、缓存命中率；')
    add_para_zh(doc, '安全测试：管理员权限控制、JWT 认证、SQL 注入与 XSS 防御。')

    add_heading_zh(doc, '1.4 测试策略', level=2)
    add_para_zh(doc, '单元测试：针对 service、router、utils 等模块的内部逻辑进行白盒测试；')
    add_para_zh(doc, '集成测试：验证 ETL 数据抓取模块与数据库、缓存（Redis）之间的数据一致性；')
    add_para_zh(doc, '系统测试：基于黑盒测试方法，从用户视角验证端到端业务流程；')
    add_para_zh(doc, '回归测试：在缺陷修复或需求变更后，重新执行相关测试用例，确保未引入新问题。')

    add_heading_zh(doc, '1.5 测试环境', level=2)
    add_para_zh(doc, '硬件环境：CPU Intel i7 / 16GB 内存 / 256GB SSD')
    add_para_zh(doc, '软件环境：Windows 11 / Python 3.12 / PostgreSQL 14 / Redis 6 / Node.js 18+')
    add_para_zh(doc, '网络环境：局域网 100Mbps，可访问互联网行情接口')

    add_heading_zh(doc, '1.6 测试进度安排', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    fill_table_header(table, ['阶段', '任务', '时间安排', '负责人'])
    stages = [
        ['测试计划', '制定测试策略、范围、资源', '第 1 周', '测试工程师'],
        ['测试设计', '编写白盒/黑盒测试用例', '第 2 周', '测试工程师'],
        ['测试执行', '运行用例、记录缺陷', '第 3-4 周', '测试工程师'],
        ['缺陷修复', '开发修复、回归验证', '第 5 周', '开发工程师'],
        ['测试报告', '汇总结果、输出报告', '第 6 周', '测试工程师'],
    ]
    for row in stages:
        table.add_row()
        fill_table_row(table, len(table.rows)-1, row)
    doc.add_paragraph()

    add_heading_zh(doc, '1.7 风险管理', level=2)
    add_para_zh(doc, '外部接口不稳定：雪球/上交所等行情源可能出现限流或返回异常，需准备 Mock 数据与降级方案；')
    add_para_zh(doc, '测试数据不足：部分策略分析依赖历史财务数据，需提前准备 seed 数据脚本；')
    add_para_zh(doc, '时间进度风险：若开发交付延期，测试时间将被压缩，需预留 20% 缓冲时间。')

    # ===== 二、白盒测试 =====
    add_heading_zh(doc, '二、白盒测试（基本路径测试）', level=1)
    add_para_zh(doc, '白盒测试基于程序内部逻辑结构，本次采用基本路径测试技术。基本路径测试的步骤为：（1）根据程序代码绘制控制流图（CFG）；（2）计算圈复杂度 V(G)；（3）确定基本路径集；（4）为每条基本路径设计测试用例。')

    # 2.1 功能一：期货数据列表查询
    add_heading_zh(doc, '2.1 功能一：期货数据列表查询（list_futures）', level=2)
    add_para_zh(doc, '2.1.1 功能描述', bold=True, indent=False)
    add_para_zh(doc, '该功能根据 symbol、frequency、limit、offset、start、end、sort 等参数，从 Redis 缓存或 PostgreSQL 数据库中查询期货数据列表，支持日频（day）和周频（week）两种粒度。若缓存命中则直接返回；否则根据 frequency 分支走周频或日频查询逻辑，最后写入缓存并返回分页结果。')

    add_para_zh(doc, '2.1.2 程序控制流图（CFG）', bold=True, indent=False)
    add_para_zh(doc, '为便于分析，将 list_futures 函数的核心判定抽象为以下控制流图节点：')
    cfg_text = """
节点说明：
① 开始：构建 cache_key，查询缓存 get_json
② 判定 D1：缓存是否有效？（isinstance(cached, dict) 且 items 为 list 且 total 为 int）
③ 判定 D2：frequency == "week"？
④ (week) 调用 _resolve_weekly_snapshot_date 获取 snapshot_date
⑤ 判定 D3：snapshot_date 是否为 None？
⑥ 判定 D4：items 为空且 snapshot_date 不为 None？（决定是否 fallback 外部抓取）
⑦ (week) 计算 total、分页 sliced、写缓存 set_json，返回结果
⑧ (day) 构建 SQLAlchemy Query
⑨ 判定 D5：symbol 是否传入？
⑩ 判定 D6：start 是否传入？
⑪ 判定 D7：end 是否传入？
⑫ 判定 D8：sort == "asc"？
⑬ (day) 执行分页查询 offset/limit、写缓存，返回结果

控制流图边（简化）：
① → ②
②(是) → 返回（路径1）
②(否) → ③
③(是) → ④ → ⑤
⑤(是) → ⑥
⑤(否) → ⑦
⑥(是) → ⑦（经 fallback）
⑥(否) → ⑦
③(否) → ⑧ → ⑨
⑨(是/否) → ⑩
⑩(是/否) → ⑪
⑪(是/否) → ⑫
⑫(是/否) → ⑬
⑬ → 返回
"""
    for line in cfg_text.strip().split('\n'):
        add_para_zh(doc, line, indent=False)

    add_para_zh(doc, '2.1.3 圈复杂度计算', bold=True, indent=False)
    add_para_zh(doc, '判定节点数 = 8（D1、D2、D3、D4、D5、D6、D7、D8）')
    add_para_zh(doc, '圈复杂度 V(G) = 判定节点数 + 1 = 8 + 1 = 9')
    add_para_zh(doc, '因此，基本路径集应包含 9 条独立路径。')

    add_para_zh(doc, '2.1.4 基本路径集', bold=True, indent=False)
    paths = [
        '路径1：①→②(是)→返回               （缓存命中，直接返回缓存数据）',
        '路径2：①→②(否)→③(是)→④→⑤(是)→⑥(是)→⑦→返回   （周频，无快照日期，fallback 后返回空或数据）',
        '路径3：①→②(否)→③(是)→④→⑤(是)→⑥(否)→⑦→返回   （周频，无快照日期，不 fallback，返回空）',
        '路径4：①→②(否)→③(是)→④→⑤(否)→⑦→返回         （周频，存在快照日期，DB 直接返回）',
        '路径5：①→②(否)→③(是)→④→⑤(否)→⑥(是)→⑦→返回   （周频，存在快照日期，但 DB 为空，fallback）',
        '路径6：①→②(否)→③(否)→⑧→⑨(是)→⑩(是)→⑪(是)→⑫(是)→⑬→返回  （日频，全条件为真，asc 排序）',
        '路径7：①→②(否)→③(否)→⑧→⑨(否)→⑩(否)→⑪(否)→⑫(否)→⑬→返回  （日频，全条件为假，desc 排序）',
        '路径8：①→②(否)→③(否)→⑧→⑨(是)→⑩(否)→⑪(是)→⑫(否)→⑬→返回  （日频，混合条件1）',
        '路径9：①→②(否)→③(否)→⑧→⑨(否)→⑩(是)→⑪(否)→⑫(是)→⑬→返回  （日频，混合条件2）',
    ]
    for p in paths:
        add_para_zh(doc, p, indent=False)

    add_para_zh(doc, '2.1.5 白盒测试用例（按表 13.1 模板）', bold=True, indent=False)
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    fill_table_header(table, ['用例编号', '功能模块', '测试标题', '前置条件', '输入数据', '预期输出', '实际输出', '测试结果'])
    wb_cases_futures = [
        ['TC-WB-01', '期货数据列表查询', '缓存命中直接返回', 'Redis 缓存中已存在有效数据', 'symbol="CU", frequency="day"', '直接返回缓存中的 items 和 total，不访问 DB', '待执行', '待执行'],
        ['TC-WB-02', '期货数据列表查询', '周频无快照日期且 fallback', 'DB 中无对应周快照记录', 'frequency="week", as_of="2099-01-01"', 'items=[], total=0（fallback 外部也无数据）', '待执行', '待执行'],
        ['TC-WB-03', '期货数据列表查询', '周频无快照日期不 fallback', 'DB 无数据，外部源也无数据', 'frequency="week", as_of="2099-01-01"', 'items=[], total=0', '待执行', '待执行'],
        ['TC-WB-04', '期货数据列表查询', '周频有快照日期直接返回', 'DB 中存在指定日期的周快照', 'frequency="week", as_of="2024-03-15"', '返回 DB 中该日期的周频数据', '待执行', '待执行'],
        ['TC-WB-05', '期货数据列表查询', '周频有快照但 DB 为空 fallback', '周快照日期存在但 DB 无记录', 'frequency="week", as_of="2024-03-15"', '通过 fallback 从外部获取数据并返回', '待执行', '待执行'],
        ['TC-WB-06', '期货数据列表查询', '日频全条件为真 asc 排序', 'symbol/start/end/sort 均传入', 'symbol="CU", start="2024-01-01", end="2024-03-01", sort="asc"', '按日期升序返回 CU 的日频数据', '待执行', '待执行'],
        ['TC-WB-07', '期货数据列表查询', '日频全条件为假 desc 排序', '仅传入 frequency', 'frequency="day"', '返回默认品种日频数据，按日期降序', '待执行', '待执行'],
        ['TC-WB-08', '期货数据列表查询', '日频混合条件1', 'symbol 传入，start/end 不传，sort=desc', 'symbol="AU", sort="desc"', '返回 AU 全部日频数据，日期降序', '待执行', '待执行'],
        ['TC-WB-09', '期货数据列表查询', '日频混合条件2', 'symbol 不传，start 传入，sort=asc', 'start="2024-01-01", sort="asc"', '返回默认品种从 start 开始的日频数据，升序', '待执行', '待执行'],
    ]
    for row in wb_cases_futures:
        table.add_row()
        fill_table_row(table, len(table.rows)-1, row)
    doc.add_paragraph()

    # 2.2 功能二：烟蒂股策略分析
    add_heading_zh(doc, '2.2 功能二：烟蒂股策略分析（analyze_cigarbutt）', level=2)
    add_para_zh(doc, '2.2.1 功能描述', bold=True, indent=False)
    add_para_zh(doc, '该功能接收股票代码 symbol，依次执行：标准化代码、从数据库构建财务面板（FinancialPanel）、获取实时行情报价、调用 CigarButtAnalyzer 进行分析、序列化指标并返回。')

    add_para_zh(doc, '2.2.2 程序控制流图（CFG）', bold=True, indent=False)
    cfg_text2 = """
节点说明：
① 开始：normalized = normalize_symbol(symbol)
② 进入 _build_financial_panel：
   查询 stock、snapshot、financial 记录
   判定 D1：stock 是否存在？（影响 stock_name、market）
   判定 D2：financial 是否存在？（影响 revenue、net_profit 等）
   判定 D3：snapshot 是否存在？（影响 total_shares、dividend_yield）
   返回 panel
③ quote = get_stock_quote(normalized)
④ 判定 D4：quote 是否非空？
   quote 非空 → stock_price = quote.get("current")
   quote 为空 → stock_price = None
⑤ analyzer = CigarButtAnalyzer(current_stock_price=stock_price)
⑥ metrics = analyzer.analyze(panel, stock_price=stock_price)
⑦ result = _metrics_to_dict(metrics)
⑧ 返回 result

控制流图边：
① → ② → ③ → ④ → ⑤ → ⑥ → ⑦ → ⑧ → 返回
"""
    for line in cfg_text2.strip().split('\n'):
        add_para_zh(doc, line, indent=False)

    add_para_zh(doc, '2.2.3 圈复杂度计算', bold=True, indent=False)
    add_para_zh(doc, '判定节点数 = 4（D1：stock 存在？ D2：financial 存在？ D3：snapshot 存在？ D4：quote 存在？）')
    add_para_zh(doc, '圈复杂度 V(G) = 判定节点数 + 1 = 4 + 1 = 5')
    add_para_zh(doc, '因此，基本路径集应包含 5 条独立路径。')

    add_para_zh(doc, '2.2.4 基本路径集', bold=True, indent=False)
    paths2 = [
        '路径1：①→②(D1 是, D2 是, D3 是)→③→④(是)→⑤→⑥→⑦→⑧→返回  （全部数据完整）',
        '路径2：①→②(D1 否, D2 否, D3 否)→③→④(否)→⑤→⑥→⑦→⑧→返回  （全部数据缺失）',
        '路径3：①→②(D1 是, D2 否, D3 是)→③→④(是)→⑤→⑥→⑦→⑧→返回  （financial 缺失）',
        '路径4：①→②(D1 是, D2 是, D3 否)→③→④(是)→⑤→⑥→⑦→⑧→返回  （snapshot 缺失）',
        '路径5：①→②(D1 否, D2 是, D3 是)→③→④(是)→⑤→⑥→⑦→⑧→返回  （stock 基础信息缺失）',
    ]
    for p in paths2:
        add_para_zh(doc, p, indent=False)

    add_para_zh(doc, '2.2.5 白盒测试用例（按表 13.1 模板）', bold=True, indent=False)
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    fill_table_header(table, ['用例编号', '功能模块', '测试标题', '前置条件', '输入数据', '预期输出', '实际输出', '测试结果'])
    wb_cases_cigar = [
        ['TC-WB-10', '烟蒂股策略分析', '全部数据完整', 'DB 中 stock/financial/snapshot 均存在，行情 API 正常', 'symbol="600519"', '返回包含完整财务指标和估值的分析结果', '待执行', '待执行'],
        ['TC-WB-11', '烟蒂股策略分析', '全部数据缺失', 'DB 中无记录，行情 API 返回空', 'symbol="UNKNOWN"', '返回基于默认值和 None 的分析结果', '待执行', '待执行'],
        ['TC-WB-12', '烟蒂股策略分析', '财务数据缺失', 'DB 中无 financial 记录', 'symbol="600519"', '分析结果中 revenue、net_profit 等字段为 None', '待执行', '待执行'],
        ['TC-WB-13', '烟蒂股策略分析', '快照数据缺失', 'DB 中无 snapshot 记录', 'symbol="600519"', '分析结果中 total_shares、dividend_yield 为 None', '待执行', '待执行'],
        ['TC-WB-14', '烟蒂股策略分析', '股票基础信息缺失', 'DB 中无 stock 记录', 'symbol="600519"', 'stock_name、market 使用默认值（normalized 码和 "A股"）', '待执行', '待执行'],
    ]
    for row in wb_cases_cigar:
        table.add_row()
        fill_table_row(table, len(table.rows)-1, row)
    doc.add_paragraph()

    # ===== 三、黑盒测试 =====
    add_heading_zh(doc, '三、黑盒测试（等价类划分与边界值分析）', level=1)
    add_para_zh(doc, '黑盒测试不考虑程序内部结构，仅从功能规格说明出发设计测试用例。本次采用等价类划分法与边界值分析法相结合的策略。')

    # 3.1 功能一：期货数据列表查询
    add_heading_zh(doc, '3.1 功能一：期货数据列表查询（list_futures_route）', level=2)
    add_para_zh(doc, '3.1.1 等价类划分', bold=True, indent=False)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    fill_table_header(table, ['输入条件', '有效等价类', '无效等价类'])
    eq_classes = [
        ['symbol', '（1）合法期货代码，如 "CU"、"AU"、"AG"', '（7）空字符串\n（8）包含特殊字符（如 @、#）\n（9）超长字符串（>20 字符）'],
        ['frequency', '（2）"day"\n（3）"week"', '（10）其他字符串（如 "month"）\n（11）空值'],
        ['limit', '（4）整数，范围 [1, 1000]', '（12）< 1\n（13）> 1000\n（14）非整数（如 10.5）'],
        ['offset', '（5）整数，范围 [0, +∞)', '（15）< 0\n（16）非整数'],
        ['start / end', '（6）合法日期格式，且 start ≤ end', '（17）非法日期格式\n（18）start > end'],
    ]
    for row in eq_classes:
        table.add_row()
        fill_table_row(table, len(table.rows)-1, row)
    doc.add_paragraph()

    add_para_zh(doc, '3.1.2 边界值分析', bold=True, indent=False)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    fill_table_header(table, ['输入条件', '边界值'])
    bv = [
        ['limit', '0, 1, 2, 999, 1000, 1001'],
        ['offset', '-1, 0, 1'],
        ['start / end', 'start = end, start = end - 1 天, end = start - 1 天'],
        ['symbol', '空字符串, 1 个字符, 6 个字符, 20 个字符, 21 个字符'],
    ]
    for row in bv:
        table.add_row()
        fill_table_row(table, len(table.rows)-1, row)
    doc.add_paragraph()

    add_para_zh(doc, '3.1.3 黑盒测试用例（按表 13.1 模板）', bold=True, indent=False)
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    fill_table_header(table, ['用例编号', '功能模块', '测试标题', '前置条件', '输入数据', '预期输出', '实际输出', '测试结果'])
    bb_cases_futures = [
        ['TC-BB-01', '期货数据列表查询', '合法 symbol + day 默认参数', '系统正常运行，DB 有数据', 'symbol="CU", frequency="day"', '返回 CU 的日频数据列表', '待执行', '待执行'],
        ['TC-BB-02', '期货数据列表查询', '合法 symbol + week', '系统正常运行，DB 有数据', 'symbol="AU", frequency="week"', '返回 AU 的周频数据列表', '待执行', '待执行'],
        ['TC-BB-03', '期货数据列表查询', 'limit 边界最小值', '系统正常运行', 'limit=1', '返回 1 条记录', '待执行', '待执行'],
        ['TC-BB-04', '期货数据列表查询', 'limit 边界最大值', '系统正常运行', 'limit=1000', '返回最多 1000 条记录', '待执行', '待执行'],
        ['TC-BB-05', '期货数据列表查询', 'offset 边界 0', '系统正常运行', 'offset=0', '从第 1 条开始返回', '待执行', '待执行'],
        ['TC-BB-06', '期货数据列表查询', '无效 frequency 值', '系统正常运行', 'frequency="month"', '返回 422 参数校验错误', '待执行', '待执行'],
        ['TC-BB-07', '期货数据列表查询', 'limit 小于下界', '系统正常运行', 'limit=0', '返回 422 参数校验错误', '待执行', '待执行'],
        ['TC-BB-08', '期货数据列表查询', 'start > end 非法区间', '系统正常运行', 'start="2024-03-01", end="2024-01-01"', '返回空结果或校验提示', '待执行', '待执行'],
        ['TC-BB-09', '期货数据列表查询', '空 symbol 走默认品种', '系统正常运行', 'symbol="", frequency="day"', '返回默认品种（CU/AU/AG 等）数据', '待执行', '待执行'],
        ['TC-BB-10', '期货数据列表查询', 'limit 大于上界', '系统正常运行', 'limit=1001', '返回 422 参数校验错误', '待执行', '待执行'],
    ]
    for row in bb_cases_futures:
        table.add_row()
        fill_table_row(table, len(table.rows)-1, row)
    doc.add_paragraph()

    # 3.2 功能二：烟蒂股策略分析
    add_heading_zh(doc, '3.2 功能二：烟蒂股策略分析（get_cigarbutt_analysis）', level=2)
    add_para_zh(doc, '3.2.1 等价类划分', bold=True, indent=False)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    fill_table_header(table, ['输入条件', '有效等价类', '无效等价类'])
    eq_classes2 = [
        ['symbol', '（1）合法 A 股代码（6 位数字）\n（2）合法港股代码（如 xxxx.HK）\n（3）合法美股代码', '（4）空字符串\n（5）包含非法字符（如 @、#）\n（6）超长字符串（>30 字符）'],
        ['数据库连接', '（7）PostgreSQL 正常连接', '（8）数据库断开或超时'],
        ['行情 API', '（9）行情服务正常响应', '（10）行情服务超时或返回异常'],
    ]
    for row in eq_classes2:
        table.add_row()
        fill_table_row(table, len(table.rows)-1, row)
    doc.add_paragraph()

    add_para_zh(doc, '3.2.2 边界值分析', bold=True, indent=False)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    fill_table_header(table, ['输入条件', '边界值'])
    bv2 = [
        ['symbol 长度', '0（空）、1、6、10、20、30、31 个字符'],
        ['symbol 格式', '纯数字、含.HK、含.、含特殊符号'],
    ]
    for row in bv2:
        table.add_row()
        fill_table_row(table, len(table.rows)-1, row)
    doc.add_paragraph()

    add_para_zh(doc, '3.2.3 黑盒测试用例（按表 13.1 模板）', bold=True, indent=False)
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    fill_table_header(table, ['用例编号', '功能模块', '测试标题', '前置条件', '输入数据', '预期输出', '实际输出', '测试结果'])
    bb_cases_cigar = [
        ['TC-BB-11', '烟蒂股策略分析', '合法 A 股代码', 'DB 有数据，行情 API 正常', 'symbol="600519"', '返回分析结果，symbol 为大写', '待执行', '待执行'],
        ['TC-BB-12', '烟蒂股策略分析', '合法港股代码', 'DB 有数据，行情 API 正常', 'symbol="0700.HK"', '返回分析结果', '待执行', '待执行'],
        ['TC-BB-13', '烟蒂股策略分析', '空字符串 symbol', '系统正常运行', 'symbol=""', '返回 400/422 参数错误', '待执行', '待执行'],
        ['TC-BB-14', '烟蒂股策略分析', '含非法字符 symbol', '系统正常运行', 'symbol="600@519"', '返回 400/422 错误或内部处理异常', '待执行', '待执行'],
        ['TC-BB-15', '烟蒂股策略分析', '超长 symbol', '系统正常运行', 'symbol="A"*50', '返回 400/422 错误或截断处理', '待执行', '待执行'],
        ['TC-BB-16', '烟蒂股策略分析', '边界长度 1', '系统正常运行', 'symbol="A"', '返回分析结果或 422 错误', '待执行', '待执行'],
        ['TC-BB-17', '烟蒂股策略分析', '数据库异常', 'PostgreSQL 断开', 'symbol="600519"', '返回 500 内部服务器错误', '待执行', '待执行'],
        ['TC-BB-18', '烟蒂股策略分析', '行情 API 超时', '行情服务不可用', 'symbol="600519"', '返回 500 错误或降级结果（stock_price=None）', '待执行', '待执行'],
    ]
    for row in bb_cases_cigar:
        table.add_row()
        fill_table_row(table, len(table.rows)-1, row)
    doc.add_paragraph()

    # 结尾
    add_para_zh(doc, '以上为 QuantPulse 系统的软件测试计划及白盒、黑盒测试用例设计。实际执行测试时，需在测试环境中逐条运行用例，记录实际输出与测试结果，并汇总缺陷报告。')

    return doc


def main():
    doc = build_document()
    docx_path = r'E:\Study\软件工程\第六次作业\第X组-QuantPulse系统-软件测试.docx'
    pdf_path = r'E:\Study\软件工程\第六次作业\第X组-QuantPulse系统-软件测试.pdf'
    doc.save(docx_path)
    print(f'Word 文档已保存：{docx_path}')

    # 尝试转换为 PDF
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        print(f'PDF 文档已保存：{pdf_path}')
    except Exception as e:
        print(f'PDF 转换失败：{e}')
        print('请手动使用 Word 打开 docx 文件并另存为 PDF。')


if __name__ == '__main__':
    main()
